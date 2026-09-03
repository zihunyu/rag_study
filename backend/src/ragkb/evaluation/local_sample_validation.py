"""Privacy-safe offline validation for authorized real format samples."""

from __future__ import annotations

import csv
import hashlib
import json
import math
import time
import tracemalloc
from collections import Counter
from collections.abc import Mapping, Sequence
from pathlib import Path
from typing import Any, cast
from urllib.parse import urlparse

import openpyxl
import xlrd
import yaml
from docx import Document as DocxDocument
from openpyxl.utils.cell import range_boundaries
from pptx import Presentation
from pypdf import PdfReader

from ragkb.application.provider_runners import embedding_provider_contract
from ragkb.config import load_env
from ragkb.contracts.provider_execution import ResultStorePort
from ragkb.document_processing.parsers import ParserRouter
from ragkb.domain.validation import DocumentQualityReport, QualityDisposition
from ragkb.engineering_security.file_validation import FORMAT_BY_EXTENSION, UploadFileValidator
from ragkb.evaluation.format_samples import _resolve

ACTIVE_FORMATS = (
    "pdf_text",
    "pdf_scanned_or_image",
    "docx",
    "pptx",
    "spreadsheet",
)
EXTERNAL_PARSER_CODES = frozenset({"OCR_REQUIRED", "MINERU_REQUIRED", "BLOCKED_EXTERNAL_PARSER"})


def _mapping(path: Path) -> dict[str, Any]:
    loaded = yaml.safe_load(path.read_text(encoding="utf-8"))
    if not isinstance(loaded, dict):
        raise ValueError("METADATA_NOT_MAPPING")
    return loaded


def _anonymous_id(format_name: str, sample: Mapping[str, object]) -> str:
    value = f"{format_name}:{sample.get('id', '')}:{sample.get('sha256', '')}"
    return hashlib.sha256(value.encode(), usedforsecurity=False).hexdigest()[:16]


def _file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        while chunk := handle.read(1024 * 1024):
            digest.update(chunk)
    return digest.hexdigest()


def _structure_counts(path: Path, source_format: str) -> dict[str, int]:
    counts: dict[str, int] = {}
    if source_format == "pdf":
        counts["pages"] = len(PdfReader(str(path)).pages)
    elif source_format == "pptx":
        counts["slides"] = len(Presentation(str(path)).slides)
    elif source_format == "docx":
        document = DocxDocument(str(path))
        counts["paragraphs"] = len(document.paragraphs)
        counts["tables"] = len(document.tables)
        counts["table_rows"] = sum(len(table.rows) for table in document.tables)
    elif source_format == "xlsx":
        xlsx_book = openpyxl.load_workbook(path, read_only=True, data_only=False)
        counts["sheets"] = len(xlsx_book.worksheets)
        counts["rows"] = sum((sheet.max_row or 0) for sheet in xlsx_book.worksheets)
        counts["columns"] = sum((sheet.max_column or 0) for sheet in xlsx_book.worksheets)
        xlsx_book.close()
    elif source_format == "xls":
        xls_book = xlrd.open_workbook(str(path), on_demand=True)
        counts["sheets"] = xls_book.nsheets
        counts["rows"] = sum(sheet.nrows for sheet in xls_book.sheets())
        counts["columns"] = sum(sheet.ncols for sheet in xls_book.sheets())
        xls_book.release_resources()
    elif source_format == "csv":
        rows = 0
        columns = 0
        with path.open("r", encoding="utf-8-sig", newline="") as handle:
            for row in csv.reader(handle):
                rows += 1
                columns = max(columns, len(row))
        counts.update(sheets=1, rows=rows, columns=columns)
    elif source_format == "image":
        counts["images"] = 1
    return counts


def _expected_locator_match(
    expected: list[Mapping[str, object]], actual: list[dict[str, object]]
) -> tuple[int, int]:
    matched = 0
    for requirement in expected:
        if "cell_range" in requirement:
            satisfied = _cell_range_fully_covered(requirement, actual)
        else:
            satisfied = any(_locator_matches(requirement, locator) for locator in actual)
        if satisfied:
            matched += 1
    return matched, len(expected)


def _cell_range_fully_covered(
    expected: Mapping[str, object], actual: list[dict[str, object]]
) -> bool:
    expected_range = expected.get("cell_range")
    if not isinstance(expected_range, str):
        return False
    try:
        bounds = range_boundaries(expected_range)
    except ValueError:
        return False
    if any(bound is None for bound in bounds):
        return False
    min_col, min_row, max_col, max_row = cast(tuple[int, int, int, int], bounds)
    expected_sheet = expected.get("sheet")
    ranges_by_row: dict[int, list[tuple[int, int]]] = {
        row: [] for row in range(min_row, max_row + 1)
    }
    for locator in actual:
        if expected_sheet is not None and locator.get("sheet") != expected_sheet:
            continue
        actual_range = locator.get("cell_range")
        if not isinstance(actual_range, str):
            continue
        try:
            actual_bounds = range_boundaries(actual_range)
        except ValueError:
            continue
        if any(bound is None for bound in actual_bounds):
            continue
        act_min_col, act_min_row, act_max_col, act_max_row = cast(
            tuple[int, int, int, int], actual_bounds
        )
        for row in range(max(min_row, act_min_row), min(max_row, act_max_row) + 1):
            ranges_by_row[row].append((max(min_col, act_min_col), min(max_col, act_max_col)))
    for intervals in ranges_by_row.values():
        cursor = min_col
        for start, end in sorted(intervals):
            if end < cursor:
                continue
            if start > cursor:
                return False
            cursor = max(cursor, end + 1)
            if cursor > max_col:
                break
        if cursor <= max_col:
            return False
    return True


def _locator_matches(expected: Mapping[str, object], actual: Mapping[str, object]) -> bool:
    for key, value in expected.items():
        name = str(key)
        if name == "row":
            if int(str(actual.get("row", 0))) != int(str(value)):
                return False
        elif actual.get(name) != value:
            return False
    return True


def _safe_error_code(error: Exception) -> str:
    code = getattr(error, "code", None)
    if isinstance(code, str) and code:
        return code
    return f"LOCAL_PARSE_ERROR_{type(error).__name__.upper()}"


def aggregate_persisted_mineru_evidence(
    samples: Sequence[Mapping[str, object]],
    results: Sequence[Mapping[str, object]],
    result_store: ResultStorePort,
) -> dict[str, object]:
    if len(samples) != len(results):
        raise ValueError("MINERU_RESULT_SAMPLE_COUNT_MISMATCH")
    completed = 0
    expected_count = 0
    matched_count = 0
    chunk_count = 0
    artifact_hashes: set[str] = set()
    for sample, result in zip(samples, results, strict=True):
        if result.get("state") != "COMPLETED":
            raise ValueError("MINERU_RESULT_NOT_COMPLETED")
        artifact_id = result.get("artifact_id")
        sample_id = result.get("sample_id")
        result_hash = result.get("result_hash")
        if (
            not isinstance(artifact_id, str)
            or not isinstance(sample_id, str)
            or not isinstance(result_hash, str)
        ):
            raise ValueError("MINERU_RESULT_EVIDENCE_INVALID")
        nodes = result_store.read_mineru_nodes(artifact_id)
        actual_locators: list[dict[str, object]] = []
        for node in nodes:
            if node.get("anonymous_sample_id") != sample_id:
                raise ValueError("MINERU_RESULT_SAMPLE_ID_MISMATCH")
            locator = node.get("locator")
            if isinstance(locator, Mapping):
                actual_locators.append(dict(locator))
            if str(node.get("display_text", "")).strip():
                chunk_count += 1
        raw_expected = sample.get("expected_locators", [])
        expected = (
            [locator for locator in raw_expected if isinstance(locator, Mapping)]
            if isinstance(raw_expected, Sequence) and not isinstance(raw_expected, (str, bytes))
            else []
        )
        matched, total = _expected_locator_match(expected, actual_locators)
        matched_count += matched
        expected_count += total
        artifact_hashes.add(result_hash)
        completed += 1
    return {
        "completed_files": completed,
        "expected_locator_count": expected_count,
        "matched_locator_count": matched_count,
        "new_chunk_count": chunk_count,
        "artifact_hash_count": len(artifact_hashes),
        "embedding_scope_unchanged_chunks": 669,
        "content_in_output": False,
        "source_names_in_output": False,
    }


def validate_local_samples(root: Path, plan_path: Path) -> dict[str, object]:
    plan = _mapping(plan_path)
    loaded = load_env(root)
    if loaded.settings is None:
        raise ValueError("CONFIG_INVALID")
    settings = loaded.settings
    validator = UploadFileValidator(max_size_bytes=settings.upload_max_file_size_mb * 1024 * 1024)
    router = ParserRouter()
    samples: list[dict[str, object]] = []
    external_by_category: Counter[str] = Counter()
    known_chunks = 0

    for item in plan.get("collection_plan", []):
        if not isinstance(item, Mapping):
            continue
        format_name = str(item.get("format", ""))
        if format_name not in ACTIVE_FORMATS or item.get("deferred_by_user"):
            continue
        directory = _resolve(root, item["sample_directory"])
        metadata = _mapping(_resolve(root, item["metadata_path"]))
        for sample in metadata.get("samples", []):
            if not isinstance(sample, Mapping):
                continue
            anonymous_id = _anonymous_id(format_name, sample)
            path = (directory / str(sample.get("file", ""))).resolve()
            status = "FAILED"
            error_code: str | None = None
            node_count = 0
            eligible_chunk_count = 0
            locator_coverage = 0.0
            disposition = "FAILED"
            issue_codes: list[str] = []
            local_route = "unknown"
            failure_stage = "preflight"
            structure: dict[str, int] = {}
            matched_locators = 0
            expected_locator_count = len(sample.get("expected_locators", []))
            before = (path.stat().st_size, path.stat().st_mtime_ns, _file_sha256(path))
            started = time.perf_counter()
            tracemalloc.start()
            try:
                extension = path.suffix.casefold()
                if extension not in FORMAT_BY_EXTENSION:
                    raise ValueError("DOC_FORMAT_UNSUPPORTED")
                source_format, declared_mime = FORMAT_BY_EXTENSION[extension]
                local_route = source_format
                failure_stage = "security_validation"
                validator.inspect(
                    path,
                    filename=path.name,
                    expected_size=path.stat().st_size,
                    expected_sha256=str(sample.get("sha256", "")),
                    declared_mime=declared_mime,
                )
                failure_stage = "structure_count"
                structure = _structure_counts(path, source_format)
                failure_stage = "parser"
                document = router.parse(source_format, path, f"sample-{anonymous_id}")
                failure_stage = "quality"
                quality = DocumentQualityReport.from_document(document)
                node_count = quality.node_count
                issue_codes = list(quality.issue_codes)
                if format_name == "pdf_scanned_or_image":
                    node_count = 0
                    status = "BLOCKED_EXTERNAL_PARSER"
                    error_code = "BLOCKED_EXTERNAL_PARSER"
                    disposition = QualityDisposition.BLOCKED_REAL_VALIDATION.value
                    external_by_category[format_name] += 1
                elif format_name == "docx":
                    status = "BLOCKED_OFFICE_CONVERSION_OR_EXTERNAL_PARSER"
                    error_code = "BLOCKED_OFFICE_CONVERSION_OR_EXTERNAL_PARSER"
                    disposition = QualityDisposition.BLOCKED_REAL_VALIDATION.value
                    issue_codes.append("OFFICE_PAGE_MAPPING_REQUIRED")
                    external_by_category[format_name] += 1
                elif quality.disposition is QualityDisposition.BLOCKED_REAL_VALIDATION:
                    status = "BLOCKED_EXTERNAL_PARSER"
                    error_code = "BLOCKED_EXTERNAL_PARSER"
                    disposition = quality.disposition.value
                    external_by_category[format_name] += 1
                else:
                    locator_coverage = quality.locator_coverage
                    disposition = quality.disposition.value
                    actual_locators = [node.locator.to_dict() for node in document.nodes]
                    failure_stage = "locator_reconciliation"
                    expected_locators = [
                        value
                        for value in sample.get("expected_locators", [])
                        if isinstance(value, Mapping)
                    ]
                    matched_locators, expected_locator_count = _expected_locator_match(
                        expected_locators, actual_locators
                    )
                    locator_gap = matched_locators < expected_locator_count
                    if locator_gap:
                        issue_codes.append("EXPECTED_LOCATOR_NOT_MATCHED")
                        disposition = QualityDisposition.DEGRADED.value
                    if locator_gap and format_name == "spreadsheet":
                        status = "BLOCKED_LOCATOR_CONTRACT"
                        error_code = "BLOCKED_LOCATOR_CONTRACT"
                    else:
                        status = (
                            "SUCCESS"
                            if quality.disposition is QualityDisposition.READY_FOR_REVIEW
                            and not locator_gap
                            else "SUCCESS_DEGRADED"
                        )
                        eligible_chunk_count = node_count
                    known_chunks += eligible_chunk_count
                    failure_stage = "complete"
            except Exception as error:  # safe aggregation: never retain exception text
                error_code = _safe_error_code(error)
                if error_code in EXTERNAL_PARSER_CODES:
                    status = "BLOCKED_EXTERNAL_PARSER"
                    external_by_category[format_name] += 1
            finally:
                _, peak_memory = tracemalloc.get_traced_memory()
                tracemalloc.stop()
            elapsed = time.perf_counter() - started
            after = (path.stat().st_size, path.stat().st_mtime_ns, _file_sha256(path))
            if before != after:
                raise RuntimeError("SAMPLE_MUTATED")
            samples.append(
                {
                    "sample_id": anonymous_id,
                    "category": format_name,
                    "local_route": local_route,
                    "status": status,
                    "error_code": error_code,
                    "failure_stage": failure_stage,
                    "node_count": node_count,
                    "chunk_count": node_count,
                    "eligible_chunk_count": eligible_chunk_count,
                    "blocked_chunk_count": node_count - eligible_chunk_count,
                    "locator_coverage": locator_coverage,
                    "expected_locator_count": expected_locator_count,
                    "matched_locator_count": matched_locators,
                    "quality_disposition": disposition,
                    "issue_codes": issue_codes,
                    "elapsed_seconds": elapsed,
                    "peak_memory_bytes": peak_memory,
                    "structure_counts": structure,
                }
            )

    by_format: dict[str, dict[str, object]] = {}
    for format_name in ACTIVE_FORMATS:
        group = [item for item in samples if item["category"] == format_name]
        status_counts = Counter(str(item["status"]) for item in group)
        structure_totals: Counter[str] = Counter()
        for sample in group:
            raw_structure = sample["structure_counts"]
            if isinstance(raw_structure, dict):
                structure_totals.update(
                    {str(key): int(str(value)) for key, value in raw_structure.items()}
                )
        by_format[format_name] = {
            "sample_count": len(group),
            "status_counts": dict(sorted(status_counts.items())),
            "node_count": sum(int(str(item["node_count"])) for item in group),
            "chunk_count": sum(int(str(item["chunk_count"])) for item in group),
            "eligible_chunk_count": sum(int(str(item["eligible_chunk_count"])) for item in group),
            "blocked_chunk_count": sum(int(str(item["blocked_chunk_count"])) for item in group),
            "locator_expected": sum(int(str(item["expected_locator_count"])) for item in group),
            "locator_matched": sum(int(str(item["matched_locator_count"])) for item in group),
            "elapsed_seconds": sum(float(str(item["elapsed_seconds"])) for item in group),
            "peak_memory_bytes_max": max(
                (int(str(item["peak_memory_bytes"])) for item in group), default=0
            ),
            "structure_counts": dict(sorted(structure_totals.items())),
        }

    external_count = sum(external_by_category.values())
    embedding_contract = embedding_provider_contract(
        base_url=settings.embedding_base_url,
        model=settings.embedding_model,
        dimension=settings.embedding_dimension,
        configured_batch_size=settings.embedding_batch_size,
        chunk_count=known_chunks,
        approved_max_batches=67,
    )
    batch_size = int(str(embedding_contract["planned_batch_size"]))
    known_batches = math.ceil(known_chunks / batch_size) if known_chunks else 0
    upper_chunks_per_file = settings.upload_max_pages * 20
    execution_failures = sum(item["status"] == "FAILED" for item in samples)
    quality_blockers = sorted(
        {str(item["error_code"]) for item in samples if str(item["status"]).startswith("BLOCKED")}
    )
    mineru_url = urlparse(settings.mineru_base_url)
    standard_mineru_v4 = bool(
        mineru_url.scheme.casefold() == "https"
        and (mineru_url.hostname or "").casefold() == "mineru.net"
        and mineru_url.path.rstrip("/") == "/api/v4"
        and not mineru_url.query
        and not mineru_url.fragment
        and mineru_url.username is None
        and mineru_url.password is None
        and mineru_url.netloc.casefold() in {"mineru.net", "mineru.net:443"}
    )
    mineru_capability = {
        "status": ("DOCX_DOCUMENTED_SUPPORTED" if standard_mineru_v4 else "CAPABILITY_UNCONFIRMED"),
        "docx_supported": standard_mineru_v4,
        "request_count": 0,
        "file_content_sent": False,
        "real_docx_execution_approved": False,
    }
    return {
        "revision": "final-local-samples:v3",
        "scope": "non_asr_5x10",
        "sample_count": len(samples),
        "execution_passed": execution_failures == 0,
        "format_quality_ready": not quality_blockers,
        "format_quality_blockers": quality_blockers,
        "by_format": by_format,
        "samples": samples,
        "embedding_budget": {
            **embedding_contract,
            "batch_size": batch_size,
            "known_chunks": known_chunks,
            "known_batches": known_batches,
            "external_chunk_formula": (
                "ceil((known_chunks + sum(external_parser_chunk_counts)) / batch_size)"
            ),
            "external_safe_upper_chunks": external_count * upper_chunks_per_file,
            "external_safe_upper_batches": math.ceil(
                (known_chunks + external_count * upper_chunks_per_file) / batch_size
            ),
        },
        "external_parser": {
            "file_count": external_count,
            "by_category": dict(sorted(external_by_category.items())),
        },
        "mineru_capability": mineru_capability,
        "external_call_count": 0,
        "network_call_performed": False,
        "sample_content_emitted": False,
        "sample_filename_emitted": False,
        "source_samples_modified": False,
        "real_acceptance": False,
    }


def external_call_plan(report: Mapping[str, object]) -> dict[str, object]:
    raw_external = report["external_parser"]
    external = raw_external if isinstance(raw_external, dict) else {}
    raw_budget = report["embedding_budget"]
    budget = raw_budget if isinstance(raw_budget, dict) else {}
    raw_categories = external.get("by_category", {})
    categories = raw_categories if isinstance(raw_categories, dict) else {}
    raw_capability = report.get("mineru_capability", {})
    capability = raw_capability if isinstance(raw_capability, dict) else {}
    return {
        "revision": "external-call-plan:v3-provider-contract-correction",
        "executed": False,
        "external_call_count": 0,
        "mineru": {
            "required_file_count": int(str(categories.get("pdf_scanned_or_image", 0))),
            "categories": {"pdf_scanned_or_image": categories.get("pdf_scanned_or_image", 0)},
            "executed": False,
            "provider_contract": "MINERU_PRECISION_API_V4",
            "current_real_request_count": 119,
            "current_state": "SCAN_AND_DOCX_PDF_COMPLETED",
            "new_retry_approved": False,
            "new_attempts": [
                {
                    "attempt_revision": "mineru-scan-attempt:v2",
                    "scope": "pdf_scanned_or_image",
                    "checkpoint_ref": "provider-checkpoints/mineru-scan-attempt-v2.json",
                    "approved_by_user": True,
                    "planned": True,
                    "executed": True,
                    "execution_status": "FAILED_HTTP_401",
                    "request_count": 1,
                    "completed_files": 0,
                    "trace_hash_count": 1,
                    "automatic_retries": 0,
                    "expected_locator_count": 10,
                },
                {
                    "attempt_revision": "mineru-scan-attempt:v3",
                    "scope": "pdf_scanned_or_image",
                    "checkpoint_ref": "provider-checkpoints/mineru-scan-attempt-v3.json",
                    "approved_by_user": True,
                    "planned": True,
                    "executed": True,
                    "execution_status": "FAILED_TASK_AFTER_UPLOAD",
                    "request_count": 3,
                    "completed_files": 0,
                    "failed_files": 1,
                    "upload_count": 1,
                    "poll_count": 1,
                    "download_count": 0,
                    "artifact_count": 0,
                    "automatic_retries": 0,
                    "expected_locator_count": 10,
                },
                {
                    "attempt_revision": "mineru-scan-attempt:v4",
                    "scope": "pdf_scanned_or_image",
                    "checkpoint_ref": "provider-checkpoints/mineru-scan-attempt-v4.json",
                    "approved_by_user": True,
                    "planned": True,
                    "executed": True,
                    "execution_status": "PARTIAL_FAILED_PROVIDER_CODE_-60002",
                    "request_count": 23,
                    "completed_files": 4,
                    "failed_files": 1,
                    "unknown_files": 0,
                    "upload_count": 4,
                    "poll_count": 10,
                    "download_count": 4,
                    "artifact_count": 4,
                    "node_count": 75,
                    "chunk_count": 75,
                    "locator_count": 75,
                    "completed_expected_locators": 4,
                    "completed_matched_locators": 4,
                    "provider_error_code": "-60002",
                    "provider_error_category": "PROVIDER_BUSINESS_ERROR_UNCLASSIFIED",
                    "trace_hash_count": 0,
                    "automatic_retries": 0,
                    "expected_locator_count": 10,
                },
                {
                    "attempt_revision": "mineru-scan-attempt:v5",
                    "scope": "pdf_scanned_or_image_positions_5_to_10",
                    "checkpoint_ref": "provider-checkpoints/mineru-scan-attempt-v5.json",
                    "approved_by_user": True,
                    "planned": True,
                    "executed": True,
                    "execution_status": "COMPLETED",
                    "request_count": 33,
                    "completed_files": 6,
                    "failed_files": 0,
                    "unknown_files": 0,
                    "upload_count": 6,
                    "poll_count": 15,
                    "download_count": 6,
                    "artifact_count": 6,
                    "node_count": 82,
                    "chunk_count": 82,
                    "locator_count": 82,
                    "completed_expected_locators": 6,
                    "completed_matched_locators": 6,
                    "automatic_retries": 0,
                    "max_files": 6,
                    "expected_locator_count": 6,
                    "derived_png_count": 1,
                    "reuses_scan_v4_completed_files": False,
                    "combined_scan_completed_files": 10,
                    "combined_expected_locators": 10,
                    "combined_matched_locators": 10,
                    "combined_artifact_count": 10,
                    "combined_node_count": 157,
                    "combined_chunk_count": 157,
                },
                {
                    "attempt_revision": "mineru-docx-attempt:v1",
                    "scope": "docx",
                    "checkpoint_ref": "provider-checkpoints/mineru-docx-attempt-v1.json",
                    "approved_by_user": True,
                    "planned": True,
                    "executed": True,
                    "execution_status": "FAILED_CONTENT_LOCATOR_INVALID",
                    "request_count": 5,
                    "completed_files": 0,
                    "failed_files": 1,
                    "unknown_files": 0,
                    "upload_count": 1,
                    "poll_count": 2,
                    "download_count": 1,
                    "artifact_count": 0,
                    "automatic_retries": 0,
                    "expected_locator_count": 20,
                },
            ],
            "authorization": {
                "max_files": 10,
                "max_requests": 330,
                "max_polls_per_file": 30,
                "poll_interval_seconds": 10,
                "automatic_retries": 0,
                "scope": "pdf_scanned_or_image",
                "runner_review_required_before_execution": True,
            },
        },
        "office_conversion_page_mapping": {
            "required_file_count": int(str(categories.get("docx", 0))),
            "categories": {"docx": categories.get("docx", 0)},
            "provider_documented": True,
            "executed": True,
            "approved_by_user": True,
            "current_attempt_status": "FAILED_CONTENT_LOCATOR_INVALID",
            "current_request_count": 7,
            "recovery": {
                "attempt_revision": "mineru-docx-recovery:v1",
                "checkpoint_ref": "provider-checkpoints/mineru-docx-recovery-v1.json",
                "approved_by_user": True,
                "planned": True,
                "executed": True,
                "execution_status": "COMPLETED_PROVIDER_RESULT_LOCATOR_GATE_FAILED",
                "request_count": 2,
                "create_count": 0,
                "upload_count": 0,
                "poll_count": 1,
                "download_count": 1,
                "completed_files": 1,
                "artifact_count": 1,
                "node_count": 27,
                "chunk_count": 27,
                "locator_count": 27,
                "expected_locator_count": 2,
                "matched_locator_count": 1,
                "locator_gate_passed": False,
                "automatic_retries": 0,
                "source_attempt_ref": "provider-checkpoints/mineru-docx-attempt-v1.json",
            },
            "remaining_attempt": {
                "attempt_revision": "mineru-docx-attempt:v2",
                "checkpoint_ref": "provider-checkpoints/mineru-docx-attempt-v2.json",
                "positions": "2-10",
                "max_files": 9,
                "max_requests": 297,
                "expected_locator_count": 18,
                "approved_by_user": True,
                "planned": True,
                "executed": False,
                "execution_status": "BLOCKED_BY_RECOVERY_LOCATOR_GATE",
            },
            "combined_target": {
                "completed_files": 10,
                "expected_locator_count": 20,
                "new_chunks_join_embedding_669": False,
            },
            "docx_pdf_attempt": {
                "attempt_revision": "mineru-docx-pdf-attempt:v1",
                "checkpoint_ref": "provider-checkpoints/mineru-docx-pdf-attempt-v1.json",
                "provider_input_revision": "libreoffice-docx-to-pdf:v1",
                "libreoffice_version": "26.8.0.3",
                "prepared_input_count": 10,
                "prepared_page_count_min": 2,
                "prepared_page_count_max": 3,
                "prepared_page_count_total": 25,
                "derived_hash_count": 10,
                "expected_pages_covered_count": 10,
                "max_files": 10,
                "max_requests": 330,
                "expected_locator_count": 20,
                "locator_policy": "strict_page_bbox",
                "approved_by_user": True,
                "planned": True,
                "executed": True,
                "execution_status": "COMPLETED",
                "request_count": 51,
                "create_count": 10,
                "upload_count": 10,
                "poll_count": 21,
                "download_count": 10,
                "completed_files": 10,
                "failed_files": 0,
                "unknown_files": 0,
                "artifact_count": 10,
                "node_count": 302,
                "chunk_count": 302,
                "locator_count": 302,
                "matched_locator_count": 20,
                "automatic_retries": 0,
                "new_chunks_join_embedding_669": False,
            },
            "capability_probe": {
                "max_read_only_requests": 0,
                "file_content_sent": False,
                "status": capability.get("status", "CAPABILITY_UNCONFIRMED"),
                "docx_supported": bool(capability.get("docx_supported", False)),
                "request_count": 0,
            },
            "native_docx_recovery_v2_deferred": True,
        },
        "embedding": {
            **budget,
            "executed": True,
            "attempt_revision": "embedding-real-attempt:v2-dashscope-batch10",
            "prior_failed_attempt_ref": "provider-checkpoints/embedding.json",
            "future_checkpoint_ref": "provider-checkpoints/embedding-attempt-v2.json",
            "required_new_batches": 67,
            "approved": True,
            "current_real_request_count": 114,
            "current_completed_batches": 113,
            "current_vector_count": 1128,
            "execution_status": "COMPLETED",
            "authorization": {
                "max_chunks": 669,
                "max_batches": 67,
                "batch_size": 10,
                "automatic_retries": 0,
                "consumed": True,
            },
            "format_remainder_attempt": {
                "attempt_revision": "embedding-real-attempt:v3-format-remainder",
                "checkpoint_ref": (
                    "provider-checkpoints/embedding-format-remainder-attempt-v3.json"
                ),
                "chunk_count": 459,
                "batch_size": 10,
                "max_batches": 46,
                "automatic_retries": 0,
                "approved_by_user": True,
                "runner_review_required_before_execution": False,
                "approved": True,
                "executed": True,
                "execution_status": "COMPLETED",
                "completed_batches": 46,
                "vector_count": 459,
                "zilliz_write_approved": False,
                "reuses_embedding_v2_checkpoint": False,
            },
            "zilliz_write_approved": False,
        },
        "reranker": {
            "executed": True,
            "checkpoint_ref": "provider-checkpoints/uat-reranker-v1.json",
            "bundle_count": 78,
            "documents_per_bundle": 4,
            "max_requests": 78,
            "request_count": 2,
            "completed_count": 1,
            "failed_count": 1,
            "unknown_count": 0,
            "gate_passed_count": 1,
            "automatic_retries": 0,
            "positive_top_k": 2,
            "execution_status": "FAILED_POSITIVE_NOT_IN_TOP_K",
            "error_code": "UAT_RERANKER_POSITIVE_NOT_IN_TOP_K",
            "global_failure_policy": "STOP_ALL_AND_DO_NOT_START_LLM",
            "conditional_user_authorization_satisfied": True,
            "runner_review_required": False,
            "second_execution_performed": False,
            "failure_review_ref": "uat-result-review/reranker-failure-1.json",
            "failure_review_sha256": (
                "8a330b506ce9b3d4d03ce99c5cf1420c06d346524c0e1ce6bda8f50c8d9eafd3"
            ),
            "provider_order_unavailable": True,
            "positive_rank_unknown": True,
            "diagnostic_v2": {
                "attempt_revision": "uat-reranker-diagnostic-runner:v2",
                "checkpoint_ref": "provider-checkpoints/uat-reranker-v2.json",
                "prior_failed_checkpoint_read_only": True,
                "max_requests": 1,
                "automatic_retries": 0,
                "planned": True,
                "approved_by_user": True,
                "executed": True,
                "runner_review_required": False,
                "request_count": 1,
                "completed_count": 1,
                "gate_passed": True,
                "positive_rank": 1,
                "response_index_count": 4,
                "llm_request_count": 0,
                "execution_status": "COMPLETED_GATE_PASSED",
            },
            "continuation_v3": {
                "attempt_revision": "uat-reranker-continuation-runner:v3",
                "checkpoint_ref": "provider-checkpoints/uat-reranker-v3.json",
                "selected_prior_results": {"v1": 1, "v2": 1},
                "remaining_candidate_count": 76,
                "max_requests": 76,
                "positive_top_k": 2,
                "automatic_retries": 0,
                "approved_by_user": True,
                "runner_review_required": False,
                "executed": True,
                "request_count": 2,
                "completed_count": 1,
                "failed_count": 1,
                "unknown_count": 0,
                "gate_passed_count": 1,
                "execution_status": "PARTIAL_GATE_FAILED",
                "error_code": "UAT_RERANKER_V3_POSITIVE_NOT_IN_TOP_K",
                "second_execution_performed": False,
            },
            "systematic_revision_v4": {
                "review_ref": "uat-systematic-revision-v4/approved-review.json",
                "manifest_ref": "uat-systematic-revision-v4/manifest.json",
                "passed_existing_count": 3,
                "pending_revision_count": 75,
                "checkpoint_ref": "provider-checkpoints/uat-reranker-v4.json",
                "max_requests": 75,
                "positive_top_k": 2,
                "automatic_retries": 0,
                "approved_by_user": False,
                "runner_review_required": True,
                "executed": True,
                "request_count": 37,
                "completed_count": 36,
                "failed_count": 1,
                "unknown_count": 0,
                "gate_passed_count": 36,
                "execution_status": "PARTIAL_GATE_FAILED",
                "error_code": "UAT_RERANKER_V4_POSITIVE_NOT_IN_TOP_K",
                "status": "EXECUTED_PARTIAL_GATE_FAILED",
            },
            "systematic_revision_v5": {
                "review_ref": "uat-systematic-revision-v5/approved-review.json",
                "manifest_ref": "uat-systematic-revision-v5/manifest.json",
                "passed_existing_count": 39,
                "pending_revision_count": 39,
                "checkpoint_ref": "provider-checkpoints/uat-reranker-v5.json",
                "max_requests": 39,
                "positive_top_k": 2,
                "automatic_retries": 0,
                "approved_by_user": False,
                "runner_review_required": True,
                "executed": False,
                "status": "PENDING_USER_REVIEW",
            },
        },
        "llm": {
            "executed": False,
            "checkpoint_ref": "provider-checkpoints/uat-llm-v1.json",
            "max_requests": 78,
            "request_count": 0,
            "completed_count": 0,
            "automatic_retries": 0,
            "execution_status": "NOT_STARTED_RERANKER_GLOBAL_GATE_FAILED",
            "prerequisite": "ALL_78_RERANKER_COMPLETED_AND_GATE_PASSED",
            "citation_gate_required": True,
            "user_result_review_required": True,
            "conditional_user_authorization_satisfied": True,
            "runner_review_required": False,
            "conditional_v2": {
                "checkpoint_ref": "provider-checkpoints/uat-llm-v2.json",
                "result_ref": "uat-results/v2",
                "candidate_count": 78,
                "max_requests": 78,
                "automatic_retries": 0,
                "prerequisite": "COMBINED_RERANKER_GATE_78_OF_78",
                "approved_by_user": True,
                "runner_review_required": False,
                "executed": False,
                "execution_status": "NOT_STARTED_RERANKER_V3_GATE_FAILED",
                "user_result_review_required": True,
                "approved_for_systematic_revision_v4": True,
                "approved_for_systematic_revision_v5": False,
            },
        },
        "uat_candidates": {
            "local_generation_approved": True,
            "status": "APPROVED_BY_USER",
            "candidate_count": 78,
            "pending_ref": "uat-candidates/pending-review.json",
            "approved_ref": "uat-candidates/approved.json",
            "approval_manifest_ref": "uat-candidates/approval-manifest.json",
            "pending_snapshot_unchanged": True,
            "model_call_count": 42,
            "reranker_call_count": 42,
            "llm_call_count": 0,
            "query_embedding_request_count": 0,
            "zilliz_request_count": 0,
        },
        "real_uat_total_model_request_budget": 156,
        "real_uat_consumed_model_requests": 42,
        "real_uat_execution_status": (
            "SYSTEMATIC_REVISION_V5_PENDING_USER_REVIEW_RERANKER_AND_LLM_NOT_APPROVED"
        ),
        "real_acceptance": False,
    }


def render_safe_summary(report: Mapping[str, object]) -> str:
    safe = {key: value for key, value in report.items() if key != "samples"}
    return json.dumps(safe, ensure_ascii=False, indent=2, sort_keys=True)
